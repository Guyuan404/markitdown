import os
import random
import string
from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont
import json
import zipfile
from datetime import datetime

class TestFileGenerator:
    def __init__(self, output_dir="test_files"):
        """初始化测试文件生成器"""
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    def _generate_random_text(self, paragraphs=3):
        """生成随机文本内容"""
        text = []
        for _ in range(paragraphs):
            words = ''.join(random.choices(string.ascii_letters + ' ' * 5, k=100))
            text.append(words)
        return '\n\n'.join(text)

    def generate_txt(self, filename="test.txt"):
        """生成文本文件"""
        path = os.path.join(self.output_dir, filename)
        content = self._generate_random_text()
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return path

    def generate_docx(self, filename="test.docx"):
        """生成 Word 文档"""
        path = os.path.join(self.output_dir, filename)
        doc = Document()
        doc.add_heading('Test Document', 0)
        
        # 添加段落
        doc.add_paragraph(self._generate_random_text())
        
        # 添加表格
        table = doc.add_table(rows=3, cols=3)
        for row in table.rows:
            for cell in row.cells:
                cell.text = ''.join(random.choices(string.ascii_letters, k=5))
        
        doc.save(path)
        return path

    def generate_xlsx(self, filename="test.xlsx"):
        """生成 Excel 文件"""
        path = os.path.join(self.output_dir, filename)
        wb = Workbook()
        ws = wb.active
        
        # 添加表头
        headers = ['ID', 'Name', 'Value']
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)
        
        # 添加数据
        for row in range(2, 6):
            ws.cell(row=row, column=1, value=row-1)
            ws.cell(row=row, column=2, value=''.join(random.choices(string.ascii_letters, k=8)))
            ws.cell(row=row, column=3, value=random.randint(1, 1000))
        
        wb.save(path)
        return path

    def generate_image(self, filename="test.png", format="PNG"):
        """生成测试图片"""
        path = os.path.join(self.output_dir, filename)
        # 创建一个300x200的图片，带有随机文本
        img = Image.new('RGB', (300, 200), color='white')
        d = ImageDraw.Draw(img)
        
        # 添加一些形状和文本
        d.rectangle([20, 20, 280, 180], outline='black')
        d.text((50, 100), f"Test Image {datetime.now()}", fill='black')
        
        img.save(path, format=format)
        return path

    def generate_html(self, filename="test.html"):
        """生成 HTML 文件"""
        path = os.path.join(self.output_dir, filename)
        content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Test HTML</title>
        </head>
        <body>
            <h1>Test Document</h1>
            <p>{self._generate_random_text()}</p>
            <table border="1">
                <tr>
                    <th>Header 1</th>
                    <th>Header 2</th>
                </tr>
                <tr>
                    <td>Data 1</td>
                    <td>Data 2</td>
                </tr>
            </table>
        </body>
        </html>
        """
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return path

    def generate_all_types(self):
        """生成所有类型的测试文件"""
        files = []
        
        # 生成文本文件
        files.append(self.generate_txt("sample.txt"))
        
        # 生成Word文档
        files.append(self.generate_docx("sample.docx"))
        
        # 生成Excel文件
        files.append(self.generate_xlsx("sample.xlsx"))
        
        # 生成HTML文件
        files.append(self.generate_html("sample.html"))
        
        # 生成不同格式的图片
        files.append(self.generate_image("sample.png", "PNG"))
        files.append(self.generate_image("sample.jpg", "JPEG"))
        
        return files

    def create_test_zip(self, filename="test_files.zip"):
        """创建包含所有测试文件的ZIP文件"""
        # 首先生成所有测试文件
        files = self.generate_all_types()
        
        # 创建ZIP文件
        zip_path = os.path.join(self.output_dir, filename)
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file in files:
                zipf.write(file, os.path.basename(file))
        
        return zip_path

def main():
    # 创建测试文件生成器
    generator = TestFileGenerator("test_files")
    
    # 生成单个测试文件
    print("Generating individual test files...")
    generator.generate_txt()
    generator.generate_docx()
    generator.generate_xlsx()
    generator.generate_html()
    generator.generate_image()
    
    # 创建包含所有测试文件的ZIP
    print("Creating test ZIP file...")
    zip_path = generator.create_test_zip()
    print(f"Test ZIP file created at: {zip_path}")

if __name__ == "__main__":
    main()
